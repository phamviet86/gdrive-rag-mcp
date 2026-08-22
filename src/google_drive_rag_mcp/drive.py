from __future__ import annotations

import hashlib
import io
import json
import os
import tempfile
from collections.abc import Iterator
from pathlib import Path
from typing import Any, cast

from docx import Document
from google.auth.exceptions import RefreshError
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from googleapiclient.http import MediaIoBaseDownload
from httplib2 import HttpLib2Error
from openpyxl import load_workbook
from pypdf import PdfReader

from .config import Settings
from .logging_utils import emit_stderr
from .models import DriveChangeBatch, SourceDocument

SCOPES = ["https://www.googleapis.com/auth/drive.readonly"]
FOLDER_MIME = "application/vnd.google-apps.folder"
GOOGLE_DOC = "application/vnd.google-apps.document"
GOOGLE_SHEET = "application/vnd.google-apps.spreadsheet"
PDF = "application/pdf"
DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
TEXT_MIMES = {"text/plain", "text/markdown", "text/x-markdown"}
SUPPORTED_MIMES = {GOOGLE_DOC, GOOGLE_SHEET, PDF, DOCX, *TEXT_MIMES}


class DriveAPIError(RuntimeError):
    def __init__(self, operation: str, status: int, reasons: tuple[str, ...]) -> None:
        self.operation = operation
        self.status = status
        self.reasons = reasons
        reason_text = ",".join(reasons) if reasons else "unknown"
        super().__init__(
            f"Google Drive API operation {operation!r} failed "
            f"with status {status or 'unknown'} and reason(s) {reason_text}"
        )


class IncompleteSearchError(RuntimeError):
    """Raised when Drive reports that a list response omitted results."""


class DownloadRestrictedError(RuntimeError):
    """Raised when Drive metadata says content cannot be downloaded or exported."""


class DriveTransportError(RuntimeError):
    """Raised after the Google SDK exhausts retries for a transport failure."""


def _http_error_details(error: HttpError) -> tuple[int, tuple[str, ...]]:
    status = int(getattr(error.resp, "status", 0) or 0)
    reasons: set[str] = set()
    try:
        raw = error.content.decode("utf-8") if isinstance(error.content, bytes) else error.content
        payload = json.loads(raw)
        details = payload.get("error", {}).get("errors", [])
        reasons.update(
            str(detail["reason"])
            for detail in details
            if isinstance(detail, dict) and detail.get("reason")
        )
    except (AttributeError, TypeError, UnicodeDecodeError, json.JSONDecodeError):
        pass
    return status, tuple(sorted(reasons))


def _is_not_found(error: HttpError) -> bool:
    status, reasons = _http_error_details(error)
    return status == 404 or "notFound" in reasons


def _secure_token_directory(path: Path) -> None:
    directory = path.expanduser().parent
    directory.mkdir(parents=True, exist_ok=True, mode=0o700)
    if os.name != "nt":
        try:
            directory.chmod(0o700)
        except OSError as error:
            raise ValueError(
                f"Cannot restrict OAuth token directory permissions at {directory}"
            ) from error


def _write_credentials_atomic(path: Path, credentials: Any) -> None:
    path = path.expanduser()
    _secure_token_directory(path)
    descriptor, temporary_name = tempfile.mkstemp(prefix=f".{path.name}.", dir=path.parent)
    temporary_path = Path(temporary_name)
    try:
        try:
            os.fchmod(descriptor, 0o600)
        except OSError:
            os.close(descriptor)
            raise
        with os.fdopen(descriptor, "w", encoding="utf-8") as stream:
            stream.write(str(credentials.to_json()))
            stream.flush()
            os.fsync(stream.fileno())
        os.replace(temporary_path, path)
        path.chmod(0o600)
        if os.name != "nt":
            directory_descriptor = os.open(path.parent, os.O_RDONLY)
            try:
                os.fsync(directory_descriptor)
            finally:
                os.close(directory_descriptor)
    finally:
        temporary_path.unlink(missing_ok=True)


def _validate_client_secret_file(path: Path) -> None:
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as error:
        raise ValueError(
            f"Cannot read Google OAuth Desktop client JSON at {path}: {error}"
        ) from error
    installed = payload.get("installed") if isinstance(payload, dict) else None
    if (
        not isinstance(installed, dict)
        or not installed.get("client_id")
        or not installed.get("client_secret")
    ):
        raise ValueError(
            "--client-secret must point to a Google OAuth Desktop client JSON containing "
            "'installed', 'client_id', and 'client_secret'"
        )


def google_credentials(
    settings: Settings,
    interactive: bool = False,
    client_secret_file: Path | None = None,
) -> Any:
    token_file = settings.token_file.expanduser()
    _secure_token_directory(token_file)
    credentials: Credentials | None = None
    if token_file.exists():
        credentials = Credentials.from_authorized_user_file(  # type: ignore[no-untyped-call]
            str(token_file), SCOPES
        )
        token_file.chmod(0o600)
        if not credentials.refresh_token:
            raise ValueError(f"OAuth token at {token_file} does not contain a refresh token")
    if credentials and not credentials.valid:
        try:
            credentials.refresh(Request())  # type: ignore[no-untyped-call]
        except RefreshError as error:
            raise ValueError(f"OAuth token refresh failed for {token_file}") from error
        if not credentials.valid:
            raise ValueError(
                f"OAuth token refresh did not produce valid credentials for {token_file}"
            )
        _write_credentials_atomic(token_file, credentials)
        emit_stderr("oauth_token_refreshed", level="info", token_path=str(token_file))
    if credentials and credentials.valid:
        return credentials
    if not interactive or client_secret_file is None:
        raise ValueError(
            "Valid Google OAuth credentials not found. Run "
            "`google-drive-rag-mcp-auth --client-secret /path/to/client_secret.json`."
        )
    _validate_client_secret_file(client_secret_file)
    flow = InstalledAppFlow.from_client_secrets_file(str(client_secret_file), SCOPES)
    credentials = flow.run_local_server(port=0)
    if not credentials.refresh_token:
        raise ValueError("Google OAuth did not return a refresh token; revoke access and retry")
    if not credentials.valid:
        raise ValueError("Google OAuth did not return valid credentials")
    _write_credentials_atomic(token_file, credentials)
    emit_stderr("oauth_token_saved", level="info", token_path=str(token_file))
    return credentials


class GoogleDriveSource:
    def __init__(self, settings: Settings) -> None:
        self.settings = settings
        self.service = build(
            "drive", "v3", credentials=google_credentials(settings), cache_discovery=False
        )
        self._closed = False

    def close(self) -> None:
        if getattr(self, "_closed", False):
            return
        self._closed = True
        close = getattr(self.service, "close", None)
        if callable(close):
            close()

    def __enter__(self) -> GoogleDriveSource:
        return self

    def __exit__(self, *_: object) -> None:
        self.close()

    def _execute(
        self,
        request: Any,
        operation: str,
        *,
        allow_not_found: bool = False,
    ) -> Any | None:
        try:
            return request.execute(num_retries=self.settings.drive_api_num_retries)
        except HttpError as error:
            if allow_not_found and _is_not_found(error):
                return None
            status, reasons = _http_error_details(error)
            emit_stderr(
                "google_drive_api_error",
                operation=operation,
                status=status,
                reasons=list(reasons),
            )
            raise DriveAPIError(operation, status, reasons) from error
        except HttpLib2Error as error:
            emit_stderr(
                "google_drive_transport_error",
                operation=operation,
                error_type=type(error).__name__,
            )
            raise DriveTransportError(
                f"Google Drive transport failed after retries during {operation!r}"
            ) from error

    def _download(self, request: Any, operation: str) -> bytes:
        output = io.BytesIO()
        downloader = MediaIoBaseDownload(
            output,
            request,
            chunksize=self.settings.drive_download_chunk_size,
        )
        done = False
        try:
            while not done:
                _, done = downloader.next_chunk(num_retries=self.settings.drive_api_num_retries)
        except HttpError as error:
            status, reasons = _http_error_details(error)
            emit_stderr(
                "google_drive_download_error",
                operation=operation,
                status=status,
                reasons=list(reasons),
            )
            raise DriveAPIError(operation, status, reasons) from error
        except HttpLib2Error as error:
            emit_stderr(
                "google_drive_transport_error",
                operation=operation,
                error_type=type(error).__name__,
            )
            raise DriveTransportError(
                f"Google Drive download transport failed after retries during {operation!r}"
            ) from error
        return output.getvalue()

    def _list_children(self, folder_id: str) -> list[dict[str, Any]]:
        result: list[dict[str, Any]] = []
        page_token: str | None = None
        while True:
            kwargs: dict[str, Any] = {
                "q": f"'{folder_id}' in parents and trashed = false",
                "fields": (
                    "nextPageToken,incompleteSearch,files(id,name,mimeType,modifiedTime,"
                    "md5Checksum,webViewLink,size,capabilities(canDownload))"
                ),
                "pageSize": 1000,
                "supportsAllDrives": True,
                "includeItemsFromAllDrives": True,
                "pageToken": page_token,
            }
            if self.settings.shared_drive_id:
                kwargs.update(
                    {
                        "corpora": "drive",
                        "driveId": self.settings.shared_drive_id,
                    }
                )
            response = cast(
                dict[str, Any],
                self._execute(self.service.files().list(**kwargs), "files.list"),
            )
            if response.get("incompleteSearch"):
                emit_stderr("google_drive_incomplete_search", operation="files.list")
                raise IncompleteSearchError(
                    "Google Drive returned incompleteSearch=true; refusing authoritative full sync"
                )
            result.extend(response.get("files", []))
            page_token = response.get("nextPageToken")
            if not page_token:
                return result

    def _metadata(self) -> Iterator[dict[str, Any]]:
        pending: list[tuple[str, tuple[str, ...], tuple[str, ...]]] = [
            (self.settings.folder_id, (), (self.settings.folder_id,))
        ]
        seen: set[str] = set()
        while pending:
            folder_id, folder_path, ancestor_folder_ids = pending.pop()
            if folder_id in seen:
                continue
            seen.add(folder_id)
            for item in self._list_children(folder_id):
                if item["mimeType"] == FOLDER_MIME:
                    pending.append(
                        (
                            item["id"],
                            (*folder_path, item["name"]),
                            (*ancestor_folder_ids, item["id"]),
                        )
                    )
                elif item["mimeType"] in SUPPORTED_MIMES:
                    item["relativePath"] = "/".join((*folder_path, item["name"]))
                    item["parentFolderId"] = folder_id
                    item["ancestorFolderIds"] = ancestor_folder_ids
                    yield item

    def _bytes(self, file_id: str, mime_type: str) -> bytes:
        if mime_type == GOOGLE_DOC:
            return self._download(
                self.service.files().export_media(fileId=file_id, mimeType="text/plain"),
                "files.export",
            )
        if mime_type == GOOGLE_SHEET:
            export_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            return self._download(
                self.service.files().export_media(fileId=file_id, mimeType=export_type),
                "files.export",
            )
        return self._download(
            self.service.files().get_media(fileId=file_id, supportsAllDrives=True),
            "files.get_media",
        )

    def _source_document(self, item: dict[str, Any]) -> SourceDocument:
        capabilities = item.get("capabilities")
        if not isinstance(capabilities, dict) or capabilities.get("canDownload") is not True:
            emit_stderr(
                "google_drive_download_restricted",
                file_id=str(item.get("id", "")),
            )
            raise DownloadRestrictedError(
                f"Google Drive does not allow downloading file {item.get('id', '')!r}"
            )
        content = self._bytes(item["id"], item["mimeType"])
        checksum = (
            item.get("md5Checksum")
            or hashlib.sha256(
                (item.get("modifiedTime", "") + str(item.get("size", ""))).encode()
            ).hexdigest()
        )
        return SourceDocument(
            id=item["id"],
            name=item["name"],
            mime_type=item["mimeType"],
            modified_time=item.get("modifiedTime", ""),
            checksum=checksum,
            web_url=item.get("webViewLink") or f"https://drive.google.com/open?id={item['id']}",
            text=self._extract(content, item["mimeType"]),
            relative_path=item["relativePath"],
            parent_folder_id=item["parentFolderId"],
            ancestor_folder_ids=tuple(item["ancestorFolderIds"]),
        )

    @staticmethod
    def _extract(content: bytes, mime_type: str) -> str:
        if mime_type in TEXT_MIMES or mime_type == GOOGLE_DOC:
            return content.decode("utf-8", errors="replace")
        if mime_type == PDF:
            return "\n\n".join(
                page.extract_text() or "" for page in PdfReader(io.BytesIO(content)).pages
            )
        if mime_type == DOCX:
            document = Document(io.BytesIO(content))
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            tables = [
                "\n".join("\t".join(cell.text for cell in row.cells) for row in table.rows)
                for table in document.tables
            ]
            return "\n\n".join([*paragraphs, *tables])
        if mime_type == GOOGLE_SHEET:
            workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
            sections = []
            for sheet in workbook.worksheets:
                rows = [
                    "\t".join("" if value is None else str(value) for value in row)
                    for row in sheet.iter_rows(values_only=True)
                ]
                sections.append(f"# Sheet: {sheet.title}\n" + "\n".join(rows))
            return "\n\n".join(sections)
        raise ValueError(f"Unsupported MIME type: {mime_type}")

    def documents(self) -> Iterator[SourceDocument]:
        for item in self._metadata():
            yield self._source_document(item)

    def start_page_token(self) -> str:
        kwargs: dict[str, Any] = {"supportsAllDrives": True}
        if self.settings.shared_drive_id:
            kwargs["driveId"] = self.settings.shared_drive_id
        response = cast(
            dict[str, Any],
            self._execute(
                self.service.changes().getStartPageToken(**kwargs),
                "changes.getStartPageToken",
            ),
        )
        return str(response["startPageToken"])

    def _get_item(self, file_id: str) -> dict[str, Any] | None:
        item = self._execute(
            self.service.files().get(
                fileId=file_id,
                fields=(
                    "id,name,mimeType,modifiedTime,md5Checksum,webViewLink,size,parents,trashed,"
                    "capabilities(canDownload)"
                ),
                supportsAllDrives=True,
            ),
            "files.get",
            allow_not_found=True,
        )
        return cast(dict[str, Any], item) if item is not None else None

    def _relative_folder_path(
        self, item: dict[str, Any]
    ) -> tuple[tuple[str, ...], tuple[str, ...]] | None:
        parents = list(item.get("parents", []))
        if not parents:
            return None
        current_id = str(parents[0])
        names: list[str] = []
        folder_ids: list[str] = []
        seen: set[str] = set()
        while current_id != self.settings.folder_id:
            if current_id in seen:
                return None
            seen.add(current_id)
            parent = self._get_item(current_id)
            if parent is None or parent.get("mimeType") != FOLDER_MIME:
                return None
            names.append(str(parent["name"]))
            folder_ids.append(current_id)
            parent_ids = list(parent.get("parents", []))
            if not parent_ids:
                return None
            current_id = str(parent_ids[0])
        names.reverse()
        folder_ids.reverse()
        return tuple(names), (self.settings.folder_id, *folder_ids)

    def _resolve_document_change(self, file_id: str) -> tuple[SourceDocument | None, bool]:
        item = self._get_item(file_id)
        if item is None:
            return None, True
        if item.get("trashed") or item.get("mimeType") not in SUPPORTED_MIMES:
            return None, False
        resolved_path = self._relative_folder_path(item)
        if resolved_path is None:
            return None, False
        folder_path, ancestor_folder_ids = resolved_path
        item["relativePath"] = "/".join((*folder_path, item["name"]))
        item["parentFolderId"] = str(item["parents"][0])
        item["ancestorFolderIds"] = ancestor_folder_ids
        return self._source_document(item), False

    def document_by_id(self, file_id: str) -> SourceDocument | None:
        document, _ = self._resolve_document_change(file_id)
        return document

    def changes(self, page_token: str) -> DriveChangeBatch:
        changed: dict[str, SourceDocument] = {}
        deleted: set[str] = set()
        full_rescan_required = False
        next_token: str | None = page_token
        new_start_page_token = ""
        while next_token:
            kwargs: dict[str, Any] = {
                "pageToken": next_token,
                "pageSize": 1000,
                "spaces": "drive",
                "includeRemoved": True,
                "includeItemsFromAllDrives": True,
                "supportsAllDrives": True,
                "fields": (
                    "nextPageToken,newStartPageToken,changes(fileId,removed,"
                    "file(id,mimeType,trashed))"
                ),
            }
            if self.settings.shared_drive_id:
                kwargs["driveId"] = self.settings.shared_drive_id
            response = cast(
                dict[str, Any],
                self._execute(self.service.changes().list(**kwargs), "changes.list"),
            )
            for change in response.get("changes", []):
                file_id = str(change.get("fileId", ""))
                item = change.get("file") or {}
                if not file_id:
                    continue
                if change.get("removed"):
                    deleted.add(file_id)
                    changed.pop(file_id, None)
                    continue
                if item.get("mimeType") == FOLDER_MIME:
                    full_rescan_required = True
                    continue
                document, delete_allowed = self._resolve_document_change(file_id)
                if delete_allowed:
                    deleted.add(file_id)
                    changed.pop(file_id, None)
                elif document is not None:
                    changed[file_id] = document
                    deleted.discard(file_id)
            next_token = response.get("nextPageToken")
            if not next_token:
                new_start_page_token = str(response.get("newStartPageToken", page_token))
        return DriveChangeBatch(
            changed_documents=tuple(changed.values()),
            delete_document_ids=frozenset(deleted),
            new_start_page_token=new_start_page_token or page_token,
            full_rescan_required=full_rescan_required,
        )


def run_oauth(settings: Settings, client_secret_file: Path) -> Path:
    google_credentials(settings, interactive=True, client_secret_file=client_secret_file)
    return settings.token_file.expanduser()
