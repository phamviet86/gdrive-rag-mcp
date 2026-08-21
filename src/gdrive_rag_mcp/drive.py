from __future__ import annotations

import hashlib
import io
from collections.abc import Iterator
from pathlib import Path
from typing import Any, cast

from docx import Document
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google.oauth2.service_account import Credentials as ServiceAccountCredentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from openpyxl import load_workbook
from pypdf import PdfReader

from .access import normalize_para, normalize_scope_value
from .config import Settings
from .models import DriveChangeBatch, SourceDocument

SCOPES = ["https://www.googleapis.com/auth/drive.readonly"]
FOLDER_MIME = "application/vnd.google-apps.folder"
GOOGLE_DOC = "application/vnd.google-apps.document"
GOOGLE_SHEET = "application/vnd.google-apps.spreadsheet"
PDF = "application/pdf"
DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
TEXT_MIMES = {"text/plain", "text/markdown", "text/x-markdown"}
SUPPORTED_MIMES = {GOOGLE_DOC, GOOGLE_SHEET, PDF, DOCX, *TEXT_MIMES}


def google_credentials(settings: Settings, interactive: bool = False) -> Any:
    if settings.service_account_file:
        return ServiceAccountCredentials.from_service_account_file(  # type: ignore[no-untyped-call]
            str(settings.service_account_file), scopes=SCOPES
        )
    credentials: Credentials | None = None
    if settings.oauth_token_file.exists():
        credentials = Credentials.from_authorized_user_file(  # type: ignore[no-untyped-call]
            str(settings.oauth_token_file), SCOPES
        )
    if credentials and credentials.expired and credentials.refresh_token:
        credentials.refresh(Request())  # type: ignore[no-untyped-call]
    if credentials and credentials.valid:
        return credentials
    if not interactive or not settings.oauth_client_file:
        raise ValueError(
            "Valid Google credentials not found. Configure a service account or run "
            "`gdrive-rag-mcp auth-google`."
        )
    flow = InstalledAppFlow.from_client_secrets_file(str(settings.oauth_client_file), SCOPES)
    credentials = flow.run_local_server(port=0)
    settings.oauth_token_file.parent.mkdir(parents=True, exist_ok=True)
    settings.oauth_token_file.write_text(credentials.to_json(), encoding="utf-8")
    settings.oauth_token_file.chmod(0o600)
    return credentials


class GoogleDriveSource:
    def __init__(self, settings: Settings) -> None:
        self.settings = settings
        self.scope_skipped = 0
        self.service = build(
            "drive", "v3", credentials=google_credentials(settings), cache_discovery=False
        )

    def _list_children(self, folder_id: str) -> list[dict[str, Any]]:
        result: list[dict[str, Any]] = []
        page_token: str | None = None
        while True:
            kwargs: dict[str, Any] = {
                "q": f"'{folder_id}' in parents and trashed = false",
                "fields": (
                    "nextPageToken,files(id,name,mimeType,modifiedTime,md5Checksum,"
                    "webViewLink,size)"
                ),
                "pageSize": 1000,
                "supportsAllDrives": True,
                "includeItemsFromAllDrives": True,
                "pageToken": page_token,
            }
            if self.settings.shared_drive_id:
                kwargs.update(
                    {
                        "corpora": "drive",
                        "driveId": self.settings.shared_drive_id,
                    }
                )
            response = self.service.files().list(**kwargs).execute()
            result.extend(response.get("files", []))
            page_token = response.get("nextPageToken")
            if not page_token:
                return result

    def _metadata(self) -> Iterator[dict[str, Any]]:
        pending: list[tuple[str, tuple[str, ...]]] = [(self.settings.folder_id, ())]
        seen: set[str] = set()
        while pending:
            folder_id, folder_path = pending.pop()
            if folder_id in seen:
                continue
            seen.add(folder_id)
            for item in self._list_children(folder_id):
                if item["mimeType"] == FOLDER_MIME:
                    pending.append((item["id"], (*folder_path, item["name"])))
                elif item["mimeType"] in SUPPORTED_MIMES:
                    classification = self._classify(folder_path)
                    if classification is None:
                        self.scope_skipped += 1
                        continue
                    item.update(classification)
                    item["relativePath"] = "/".join((*folder_path, item["name"]))
                    item["parentFolderId"] = folder_id
                    yield item

    def _classify(self, folder_path: tuple[str, ...]) -> dict[str, str] | None:
        if self.settings.scope_layout == "flat":
            return {
                "ownerProfileId": "shared",
                "businessFunction": "general",
                "paraCategory": "resources",
            }
        if len(folder_path) < 3:
            return None
        owner = normalize_scope_value(folder_path[0])
        business_function = normalize_scope_value(folder_path[1])
        para_category = normalize_para(folder_path[2])
        if (
            not owner
            or not business_function
            or para_category
            not in {
                "projects",
                "areas",
                "resources",
                "archives",
            }
        ):
            return None
        return {
            "ownerProfileId": owner,
            "businessFunction": business_function,
            "paraCategory": para_category,
        }

    def _bytes(self, file_id: str, mime_type: str) -> bytes:
        if mime_type == GOOGLE_DOC:
            result = (
                self.service.files().export_media(fileId=file_id, mimeType="text/plain").execute()
            )
            return cast(bytes, result)
        if mime_type == GOOGLE_SHEET:
            export_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            result = (
                self.service.files().export_media(fileId=file_id, mimeType=export_type).execute()
            )
            return cast(bytes, result)
        result = self.service.files().get_media(fileId=file_id, supportsAllDrives=True).execute()
        return cast(bytes, result)

    def _source_document(self, item: dict[str, Any]) -> SourceDocument:
        content = self._bytes(item["id"], item["mimeType"])
        checksum = (
            item.get("md5Checksum")
            or hashlib.sha256(
                (item.get("modifiedTime", "") + str(item.get("size", ""))).encode()
            ).hexdigest()
        )
        return SourceDocument(
            id=item["id"],
            name=item["name"],
            mime_type=item["mimeType"],
            modified_time=item.get("modifiedTime", ""),
            checksum=checksum,
            web_url=item.get("webViewLink") or f"https://drive.google.com/open?id={item['id']}",
            text=self._extract(content, item["mimeType"]),
            owner_profile_id=item["ownerProfileId"],
            business_function=item["businessFunction"],
            para_category=item["paraCategory"],
            relative_path=item["relativePath"],
            parent_folder_id=item["parentFolderId"],
        )

    @staticmethod
    def _extract(content: bytes, mime_type: str) -> str:
        if mime_type in TEXT_MIMES or mime_type == GOOGLE_DOC:
            return content.decode("utf-8", errors="replace")
        if mime_type == PDF:
            return "\n\n".join(
                page.extract_text() or "" for page in PdfReader(io.BytesIO(content)).pages
            )
        if mime_type == DOCX:
            document = Document(io.BytesIO(content))
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            tables = [
                "\n".join("\t".join(cell.text for cell in row.cells) for row in table.rows)
                for table in document.tables
            ]
            return "\n\n".join([*paragraphs, *tables])
        if mime_type == GOOGLE_SHEET:
            workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
            sections = []
            for sheet in workbook.worksheets:
                rows = [
                    "\t".join("" if value is None else str(value) for value in row)
                    for row in sheet.iter_rows(values_only=True)
                ]
                sections.append(f"# Sheet: {sheet.title}\n" + "\n".join(rows))
            return "\n\n".join(sections)
        raise ValueError(f"Unsupported MIME type: {mime_type}")

    def documents(self) -> Iterator[SourceDocument]:
        for item in self._metadata():
            yield self._source_document(item)

    def start_page_token(self) -> str:
        kwargs: dict[str, Any] = {"supportsAllDrives": True}
        if self.settings.shared_drive_id:
            kwargs["driveId"] = self.settings.shared_drive_id
        response = self.service.changes().getStartPageToken(**kwargs).execute()
        return str(response["startPageToken"])

    def _get_item(self, file_id: str) -> dict[str, Any] | None:
        try:
            item = (
                self.service.files()
                .get(
                    fileId=file_id,
                    fields=(
                        "id,name,mimeType,modifiedTime,md5Checksum,webViewLink,size,parents,trashed"
                    ),
                    supportsAllDrives=True,
                )
                .execute()
            )
        except HttpError as error:
            if error.resp.status in {403, 404}:
                return None
            raise
        return item if not item.get("trashed") else None

    def _relative_folder_path(self, item: dict[str, Any]) -> tuple[str, ...] | None:
        parents = list(item.get("parents", []))
        if not parents:
            return None
        current_id = str(parents[0])
        names: list[str] = []
        seen: set[str] = set()
        while current_id != self.settings.folder_id:
            if current_id in seen:
                return None
            seen.add(current_id)
            parent = self._get_item(current_id)
            if parent is None or parent.get("mimeType") != FOLDER_MIME:
                return None
            names.append(str(parent["name"]))
            parent_ids = list(parent.get("parents", []))
            if not parent_ids:
                return None
            current_id = str(parent_ids[0])
        names.reverse()
        return tuple(names)

    def document_by_id(self, file_id: str) -> SourceDocument | None:
        item = self._get_item(file_id)
        if item is None or item.get("mimeType") not in SUPPORTED_MIMES:
            return None
        folder_path = self._relative_folder_path(item)
        if folder_path is None:
            return None
        classification = self._classify(folder_path)
        if classification is None:
            self.scope_skipped += 1
            return None
        item.update(classification)
        item["relativePath"] = "/".join((*folder_path, item["name"]))
        item["parentFolderId"] = str(item["parents"][0])
        return self._source_document(item)

    def changes(self, page_token: str) -> DriveChangeBatch:
        changed: dict[str, SourceDocument] = {}
        deleted: set[str] = set()
        full_rescan_required = False
        next_token: str | None = page_token
        new_start_page_token = ""
        initial_skipped = self.scope_skipped
        while next_token:
            kwargs: dict[str, Any] = {
                "pageToken": next_token,
                "pageSize": 1000,
                "spaces": "drive",
                "includeRemoved": True,
                "includeItemsFromAllDrives": True,
                "supportsAllDrives": True,
                "fields": (
                    "nextPageToken,newStartPageToken,changes(fileId,removed,"
                    "file(id,mimeType,trashed))"
                ),
            }
            if self.settings.shared_drive_id:
                kwargs["driveId"] = self.settings.shared_drive_id
            response = self.service.changes().list(**kwargs).execute()
            for change in response.get("changes", []):
                file_id = str(change.get("fileId", ""))
                item = change.get("file") or {}
                if not file_id:
                    continue
                if change.get("removed") or item.get("trashed"):
                    deleted.add(file_id)
                    changed.pop(file_id, None)
                    continue
                if item.get("mimeType") == FOLDER_MIME:
                    full_rescan_required = True
                    continue
                document = self.document_by_id(file_id)
                if document is None:
                    deleted.add(file_id)
                    changed.pop(file_id, None)
                else:
                    changed[file_id] = document
                    deleted.discard(file_id)
            next_token = response.get("nextPageToken")
            if not next_token:
                new_start_page_token = str(response.get("newStartPageToken", page_token))
        return DriveChangeBatch(
            changed_documents=tuple(changed.values()),
            delete_document_ids=frozenset(deleted),
            new_start_page_token=new_start_page_token or page_token,
            full_rescan_required=full_rescan_required,
            scope_skipped=self.scope_skipped - initial_skipped,
        )


def run_oauth(settings: Settings) -> Path:
    google_credentials(settings, interactive=True)
    return settings.oauth_token_file
